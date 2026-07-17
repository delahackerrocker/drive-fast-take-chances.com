from pathlib import Path
import json
import re
import subprocess
import zipfile

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "tmp" / "resume" / "Steven_de_la_Torre_Master_Resume_3_Page.md"
DOCX = ROOT / "output" / "resume" / "Steven_de_la_Torre_Master_Resume_3_Page.docx"
PDF = ROOT / "output" / "resume" / "Steven_de_la_Torre_Master_Resume_3_Page.pdf"
TMP = ROOT / "tmp" / "resume"


def docx_text_and_xml():
    with zipfile.ZipFile(DOCX) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    text_nodes = re.findall(r"<w:t(?: [^>]*)?>(.*?)</w:t>", document_xml)
    text = "\n".join(re.sub(r"&amp;", "&", node) for node in text_nodes)
    return text, document_xml, rels_xml


def embedded_pdf_fonts(reader):
    fonts = set()
    for page in reader.pages:
        resources = page["/Resources"].get_object()
        for font_ref in resources.get("/Font", {}).values():
            font = font_ref.get_object()
            descriptor_ref = font.get("/FontDescriptor")
            embedded = False
            if descriptor_ref:
                descriptor = descriptor_ref.get_object()
                embedded = any(key in descriptor for key in ("/FontFile", "/FontFile2", "/FontFile3"))
            fonts.add((str(font.get("/BaseFont")), str(font.get("/Subtype")), embedded))
    return [dict(base_font=a, subtype=b, embedded=c) for a, b, c in sorted(fonts)]


def original_files_unchanged():
    originals = [
        "output/resume/Steven_de_la_Torre_Master_Resume.md",
        "output/resume/Steven_de_la_Torre_Master_Resume.docx",
        "output/resume/Steven_de_la_Torre_Master_Resume.pdf",
    ]
    result = subprocess.run(
        ["git", "diff", "--exit-code", "--", *originals],
        cwd=ROOT,
        capture_output=True,
        text=True,
    )
    return result.returncode == 0


def main():
    checks = {
        "source_exists": SOURCE.exists() and SOURCE.stat().st_size > 0,
        "docx_exists": DOCX.exists() and DOCX.stat().st_size > 0,
        "pdf_exists": PDF.exists() and PDF.stat().st_size > 0,
        "original_two_page_files_unchanged": original_files_unchanged(),
    }

    source = SOURCE.read_text(encoding="utf-8")
    checks["source_no_placeholders"] = not re.search(
        r"INSERT CURRENT|\bTBD\b|\bTODO\b|drive-fast-take-chances\.com", source, re.I
    )
    checks["portfolio_url_current"] = "https://practitioner.digital" in source

    document = Document(DOCX)
    docx_text, document_xml, rels_xml = docx_text_and_xml()
    (TMP / "docx-text-3-page.txt").write_text(docx_text, encoding="utf-8")
    section = document.sections[0]
    checks["docx_one_section"] = len(document.sections) == 1
    checks["docx_letter_page"] = (
        round(section.page_width.inches, 2) == 8.5
        and round(section.page_height.inches, 2) == 11.0
    )
    checks["docx_margins"] = (
        round(section.left_margin.inches, 2) == 0.65
        and round(section.right_margin.inches, 2) == 0.65
        and round(section.top_margin.inches, 2) == 0.58
        and round(section.bottom_margin.inches, 2) == 0.58
    )
    markdown_bullets = len(re.findall(r"(?m)^- ", source))
    numbered_paragraphs = sum(
        1
        for paragraph in document.paragraphs
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    )
    checks["docx_real_bullets"] = numbered_paragraphs == markdown_bullets and numbered_paragraphs > 0
    checks["docx_two_manual_page_breaks"] = document_xml.count("<w:pageBreakBefore") == 2
    checks["docx_hyperlinks"] = all(
        target in rels_xml
        for target in (
            "mailto:SteveDeLaTorre@gmail.com",
            "https://practitioner.digital",
            "https://www.linkedin.com/in/delahackerrocker/",
        )
    )
    checks["docx_no_textboxes_or_drawings"] = (
        "w:txbxContent" not in document_xml and "w:drawing" not in document_xml
    )
    checks["docx_empty_header_footer"] = all(
        not paragraph.text.strip()
        for sec in document.sections
        for part in (sec.header, sec.footer)
        for paragraph in part.paragraphs
    )

    reader = PdfReader(PDF)
    pdf_pages_text = [page.extract_text() or "" for page in reader.pages]
    pdf_text = "\n".join(pdf_pages_text)
    (TMP / "pdf-text-3-page.txt").write_text(pdf_text, encoding="utf-8")
    checks["pdf_three_pages"] = len(reader.pages) == 3
    checks["pdf_nonblank_pages"] = all(
        len(re.sub(r"\s+", "", page_text)) > 200 for page_text in pdf_pages_text
    )
    checks["pdf_letter_pages"] = all(
        round(float(page.mediabox.width)) == 612
        and round(float(page.mediabox.height)) == 792
        for page in reader.pages
    )
    link_annotations = sum(
        1
        for page in reader.pages
        for annotation_ref in page.get("/Annots", [])
        if annotation_ref.get_object().get("/Subtype") == "/Link"
    )
    checks["pdf_hyperlinks"] = link_annotations >= 3
    fonts = embedded_pdf_fonts(reader)
    safe_base14 = {
        "/Helvetica", "/Helvetica-Bold", "/Helvetica-Oblique", "/Helvetica-BoldOblique",
        "/Times-Roman", "/Times-Bold", "/Times-Italic", "/Times-BoldItalic",
        "/Courier", "/Courier-Bold", "/Courier-Oblique", "/Courier-BoldOblique",
        "/Symbol", "/ZapfDingbats",
    }
    checks["pdf_fonts_embedded_or_safe"] = bool(fonts) and all(
        item["embedded"] or item["base_font"] in safe_base14 for item in fonts
    )

    essential = [
        "STEVEN DE LA TORRE",
        "IMPREZARIO ENTERTAINMENT",
        "ACTIVISION BLIZZARD / RAVEN SOFTWARE",
        "PROFESSIONAL EXPERIENCE — CONTINUED",
        "SABERTOOTH",
        "YAHOO!",
        "WARNER BROS. ONLINE",
        "PRACT1T10N3R / GR1M01RE",
        "AWARDS AND RECOGNITION",
        "EDUCATION AND CERTIFICATIONS",
        "AbleGamers Academy",
    ]
    checks["docx_essential_text"] = all(item.lower() in docx_text.lower() for item in essential)
    checks["pdf_essential_text"] = all(item.lower() in pdf_text.lower() for item in essential)
    checks["contact_selectable"] = all(
        item.lower() in docx_text.lower() and item.lower() in pdf_text.lower()
        for item in ("Steven de la Torre", "SteveDeLaTorre@gmail.com", "practitioner.digital")
    )

    report = {
        "checks": checks,
        "all_passed": all(checks.values()),
        "source_words": len(source.split()),
        "source_bullets": markdown_bullets,
        "docx_numbered_paragraphs": numbered_paragraphs,
        "docx_page_breaks": document_xml.count("<w:pageBreakBefore"),
        "pdf_pages": len(reader.pages),
        "pdf_link_annotations": link_annotations,
        "pdf_fonts": fonts,
    }
    print(json.dumps(report, indent=2))
    if not report["all_passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
