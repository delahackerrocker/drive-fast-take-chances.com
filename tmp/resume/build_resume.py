from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "output" / "resume" / "Steven_de_la_Torre_Master_Resume.md"
OUTPUT = ROOT / "output" / "resume" / "Steven_de_la_Torre_Master_Resume.docx"
ROLE_PAGE_BREAKS = {
    "FREELANCE — Mixed Reality Developer": "PROFESSIONAL EXPERIENCE — CONTINUED",
}
SECTION_PAGE_BREAKS = set()

CYAN = "12B5CD"
CYAN_DARK = RGBColor(0x00, 0x78, 0x88)
BLACK = RGBColor(0x05, 0x05, 0x05)
GRAY = RGBColor(0x4F, 0x4F, 0x4F)
LIGHT_GRAY = RGBColor(0x67, 0x67, 0x67)
FONT = "Aptos"
DISPLAY_FONT = "Aptos Display"


def set_font(run, name=FONT, size=8.55, color=BLACK, bold=None, italic=None):
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_free_paragraph_geometry(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True


def set_keep(paragraph, keep_next=False, keep_lines=True):
    paragraph.paragraph_format.keep_with_next = keep_next
    paragraph.paragraph_format.keep_together = keep_lines


def add_hyperlink(paragraph, text, url, size=8.35, bold=False):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "007888")
    rpr.append(color)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    rpr.append(size_el)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(size * 2)))
    rpr.append(size_cs)
    if bold:
        rpr.append(OxmlElement("w:b"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*.*?\*\*|\[.*?\]\(.*?\))")


def add_inline(paragraph, text, size=8.55, color=BLACK):
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            set_font(paragraph.add_run(text[cursor:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, color=color, bold=True)
        else:
            link = re.match(r"\[(.*?)\]\((.*?)\)", token)
            add_hyperlink(paragraph, link.group(1), link.group(2), size=size)
        cursor = match.end()
    if cursor < len(text):
        set_font(paragraph.add_run(text[cursor:]), size=size, color=color)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def add_real_bullet_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "288")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "288")
    ind.set(qn("w:hanging"), "144")
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "007888")
    rpr.append(color)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    ppr.insert(0, num_pr)


def setup_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(8.55)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.03

    for name in ("Resume Name", "Resume Section", "Resume Role", "Resume Meta", "Resume Tech"):
        if name not in document.styles:
            document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    name_style = document.styles["Resume Name"]
    name_style.font.name = DISPLAY_FONT
    name_style._element.rPr.rFonts.set(qn("w:ascii"), DISPLAY_FONT)
    name_style._element.rPr.rFonts.set(qn("w:hAnsi"), DISPLAY_FONT)
    name_style.font.size = Pt(24)
    name_style.font.bold = True
    name_style.font.color.rgb = BLACK
    name_style.paragraph_format.space_after = Pt(0.5)
    name_style.paragraph_format.keep_with_next = True

    section_style = document.styles["Resume Section"]
    section_style.font.name = DISPLAY_FONT
    section_style._element.rPr.rFonts.set(qn("w:ascii"), DISPLAY_FONT)
    section_style._element.rPr.rFonts.set(qn("w:hAnsi"), DISPLAY_FONT)
    section_style.font.size = Pt(9.1)
    section_style.font.bold = True
    section_style.font.color.rgb = BLACK
    section_style.paragraph_format.space_before = Pt(4.5)
    section_style.paragraph_format.space_after = Pt(2.2)
    section_style.paragraph_format.line_spacing = 1.0
    section_style.paragraph_format.keep_with_next = True

    role_style = document.styles["Resume Role"]
    role_style.font.name = DISPLAY_FONT
    role_style._element.rPr.rFonts.set(qn("w:ascii"), DISPLAY_FONT)
    role_style._element.rPr.rFonts.set(qn("w:hAnsi"), DISPLAY_FONT)
    role_style.font.size = Pt(9.3)
    role_style.font.bold = True
    role_style.font.color.rgb = BLACK
    role_style.paragraph_format.space_before = Pt(3.4)
    role_style.paragraph_format.space_after = Pt(0)
    role_style.paragraph_format.line_spacing = 1.0
    role_style.paragraph_format.keep_with_next = True

    meta_style = document.styles["Resume Meta"]
    meta_style.font.name = FONT
    meta_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    meta_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    meta_style.font.size = Pt(7.9)
    meta_style.font.bold = True
    meta_style.font.color.rgb = GRAY
    meta_style.paragraph_format.space_before = Pt(0)
    meta_style.paragraph_format.space_after = Pt(0)
    meta_style.paragraph_format.line_spacing = 1.0
    meta_style.paragraph_format.keep_with_next = True

    tech_style = document.styles["Resume Tech"]
    tech_style.font.name = FONT
    tech_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    tech_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    tech_style.font.size = Pt(7.55)
    tech_style.font.color.rgb = CYAN_DARK
    tech_style.paragraph_format.space_before = Pt(0)
    tech_style.paragraph_format.space_after = Pt(1.2)
    tech_style.paragraph_format.line_spacing = 1.0
    tech_style.paragraph_format.keep_with_next = True


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    setup_styles(document)
    bullet_num_id = add_real_bullet_numbering(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    current_section = ""
    after_role_meta = False
    first_line = True

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Resume Name")
            set_font(paragraph.add_run(line[2:]), name=DISPLAY_FONT, size=24, bold=True)
            first_line = False
            continue

        if first_line:
            continue

        if line.startswith("## "):
            current_section = line[3:]
            paragraph = document.add_paragraph(style="Resume Section")
            if current_section in SECTION_PAGE_BREAKS:
                paragraph.paragraph_format.page_break_before = True
            shade_paragraph(paragraph, CYAN)
            set_font(paragraph.add_run(current_section.upper()), name=DISPLAY_FONT, size=9.1, bold=True)
            set_keep(paragraph, keep_next=True)
            after_role_meta = False
            continue

        if line.startswith("### "):
            heading = line[4:]
            continuation_label = next(
                (label for prefix, label in ROLE_PAGE_BREAKS.items() if heading.startswith(prefix)),
                None,
            )
            if continuation_label:
                continuation = document.add_paragraph(style="Resume Section")
                continuation.paragraph_format.page_break_before = True
                shade_paragraph(continuation, CYAN)
                set_font(
                    continuation.add_run(continuation_label),
                    name=DISPLAY_FONT,
                    size=9.1,
                    bold=True,
                )
                set_keep(continuation, keep_next=True)
            paragraph = document.add_paragraph(style="Resume Role")
            set_font(paragraph.add_run(heading), name=DISPLAY_FONT, size=9.3, bold=True)
            set_keep(paragraph, keep_next=True)
            after_role_meta = True
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, bullet_num_id)
            set_cell_free_paragraph_geometry(paragraph, after=1.2, line=1.0)
            set_keep(paragraph, keep_lines=True)
            add_inline(paragraph, line[2:], size=8.25)
            after_role_meta = False
            continue

        if line.startswith("**Systems Designer"):
            paragraph = document.add_paragraph()
            set_cell_free_paragraph_geometry(paragraph, after=1.8, line=1.0)
            set_keep(paragraph, keep_next=True)
            set_font(paragraph.add_run(line[2:-2]), name=DISPLAY_FONT, size=10.25, color=CYAN_DARK, bold=True)
            continue

        if line.startswith("Columbus, OH"):
            paragraph = document.add_paragraph()
            set_cell_free_paragraph_geometry(paragraph, after=0, line=1.0)
            set_keep(paragraph, keep_next=True)
            parts = [part.strip() for part in line.replace("  ", "").split("·")]
            for index, part in enumerate(parts):
                if index:
                    set_font(paragraph.add_run("  ·  "), size=8.35, color=LIGHT_GRAY)
                if "@" in part:
                    add_hyperlink(paragraph, part, f"mailto:{part}", size=8.35)
                else:
                    set_font(paragraph.add_run(part), size=8.35, color=GRAY)
            continue

        if line.startswith("[Portfolio]"):
            paragraph = document.add_paragraph()
            set_cell_free_paragraph_geometry(paragraph, after=3.0, line=1.0)
            add_hyperlink(paragraph, "practitioner.digital", "https://practitioner.digital", size=8.35, bold=True)
            set_font(paragraph.add_run("  ·  "), size=8.35, color=LIGHT_GRAY)
            add_hyperlink(
                paragraph,
                "linkedin.com/in/delahackerrocker",
                "https://www.linkedin.com/in/delahackerrocker/",
                size=8.35,
            )
            continue

        if after_role_meta and line.startswith("**"):
            paragraph = document.add_paragraph(style="Resume Meta")
            add_inline(paragraph, line.replace("  ", ""), size=7.9, color=GRAY)
            set_keep(paragraph, keep_next=True)
            after_role_meta = "await-tech"
            continue

        if after_role_meta == "await-tech":
            paragraph = document.add_paragraph(style="Resume Tech")
            set_font(paragraph.add_run(line), size=7.55, color=CYAN_DARK)
            set_keep(paragraph, keep_next=True)
            after_role_meta = False
            continue

        paragraph = document.add_paragraph()
        compact_section = current_section in {"Core Capabilities", "Skills", "Awards and Recognition", "Education and Certifications"}
        size = 7.75 if compact_section else 8.45
        after = 1.1 if compact_section else 2.2
        set_cell_free_paragraph_geometry(paragraph, after=after, line=1.0 if compact_section else 1.03)
        set_keep(paragraph, keep_lines=True)
        add_inline(paragraph, line.replace("  ", ""), size=size, color=BLACK)

    core_props = document.core_properties
    core_props.title = "Steven de la Torre Master Resume"
    core_props.subject = "Systems Design, Software Development, UX, and Interactive Design"
    core_props.author = "Steven de la Torre"
    core_props.keywords = "Systems Design, Software Development, UX, Game Development, Unity, Unreal Engine"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
